#!/usr/bin/env python3
"""Builds REPORT.docx for Project 9 — Ask Your Logs.

Simple, readable report: each step has the command(s) to run and a clearly
marked placeholder where you paste your own output or screenshot.
Run:  python3 build_report.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def para(text=""):
    doc.add_paragraph(text)


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def code(text):
    """Monospace block for commands."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    return p


def placeholder(label="PASTE YOUR OUTPUT HERE"):
    p = doc.add_paragraph()
    run = p.add_run(f"[ {label} ]")
    run.italic = True
    run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
    run.font.size = Pt(10)
    return p


# ============================= TITLE =============================
title = doc.add_heading("Project 9 — Ask Your Logs", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Monitoring, Observability, and Keeping Watch")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = doc.add_paragraph("Name: ____________________    Course: System Administration & Maintenance    Date: __________")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ============================= SUMMARY =============================
h1("1. Executive Summary")
para(
    "Cornerstone Faith Resources had a ~90-second checkout outage at 14:05. "
    "This project stands up a real observability stack (Prometheus + Grafana + "
    "OpenSearch), reproduces that incident from the logs, and adds a "
    "natural-language query layer so on-call can ASK the logs instead of "
    "grepping them by hand."
)
para("Finding, in one line:")
bullet("SYMPTOM: the checkout service aborted with reason 'payment_declined' for three users.")
bullet("ROOT CAUSE: the payments service could not authorize because the upstream provider (stripe) timed out — HTTP 504, ~5-second latency.")
bullet("PROOF: each 'checkout aborted' line shares its trace_id with a 'payments authorize failed' line (d4e1, e2a8, aa01).")
bullet("Affected users: u-4419, u-4420, u-4422. The system recovered on its own by 14:07.")

# ============================= STEP 1 =============================
h1("2. Step 1 — Bring Up the Stack (all four UIs healthy)")
para("Start the observability stack with Docker Compose:")
code("docker compose -f docker-compose.observability.yml up -d\n"
     "docker compose -f docker-compose.observability.yml ps")
para("Command output (all containers 'Up'):")
placeholder("PASTE the `docker compose ps` output here")
para()
para("Confirm all four UIs respond. In a Codespace, open each forwarded port; "
     "locally, open each URL:")
bullet("Prometheus  -> http://localhost:9090   (Status > Targets)")
bullet("Grafana     -> http://localhost:3001   (login admin / admin)")
bullet("OpenSearch  -> http://localhost:9200   (returns cluster JSON)")
bullet("OpenSearch Dashboards -> http://localhost:5601")
para("Quick health check from the terminal:")
code("curl -s http://localhost:9090/-/healthy\n"
     "curl -s http://localhost:9200 | head -20\n"
     "curl -s -o /dev/null -w '%{http_code}\\n' http://localhost:3001/login\n"
     "curl -s -o /dev/null -w '%{http_code}\\n' http://localhost:5601/api/status")
placeholder("PASTE the health-check output here")
para("SCREENSHOTS — paste one screenshot of EACH of the four UIs healthy:")
placeholder("SCREENSHOT: Prometheus")
placeholder("SCREENSHOT: Grafana")
placeholder("SCREENSHOT: OpenSearch (9200 JSON)")
placeholder("SCREENSHOT: OpenSearch Dashboards")

# ============================= STEP 2 =============================
h1("3. Step 2 — Load the Logs (verify 14 documents)")
para("Load sample-logs.jsonl into OpenSearch with the _bulk recipe and verify "
     "the document count:")
code("chmod +x load-logs.sh\n./load-logs.sh")
para("Expected: the count at the end reads 14. Paste the output:")
placeholder("PASTE the load-logs.sh output — must show \"count\" : 14")

# ============================= STEP 3 =============================
h1("4. Step 3 — Grafana: p95 Latency + Error-Ratio Panels")
para("Add Prometheus as a data source in Grafana (URL http://prometheus:9090), "
     "then build two panels. The PromQL is shown below (see grafana-queries.md).")
h2("Panel 1 — p95 latency")
code("histogram_quantile(0.95, sum(rate(prometheus_http_request_duration_seconds_bucket[5m])) by (le))")
h2("Panel 2 — error ratio")
code("sum(rate(prometheus_http_requests_total{code=~\"5..\"}[5m]))\n"
     "  /\n"
     "sum(rate(prometheus_http_requests_total[5m]))")
para("SCREENSHOT — the Grafana dashboard with BOTH panels and their PromQL visible:")
placeholder("SCREENSHOT: Grafana dashboard (p95 + error-ratio, PromQL shown)")

# ============================= STEP 4 =============================
h1("5. Step 4 — Answer the Three Incident Questions BY HAND")
para("These are exact OpenSearch queries I wrote by hand (full versions in "
     "queries-by-hand.md). Run each and paste the JSON result.")

h2("Q1. Why did checkout fail at 14:05?")
code("curl -s 'http://localhost:9200/storefront-logs/_search?pretty' \\\n"
     "  -H 'Content-Type: application/json' -d '{\n"
     "  \"size\": 20,\n"
     "  \"query\": { \"bool\": { \"filter\": [\n"
     "    { \"match\": { \"level\": \"ERROR\" } },\n"
     "    { \"range\": { \"ts\": { \"gte\": \"2026-06-12T14:05:00\", \"lt\": \"2026-06-12T14:06:00\" } } }\n"
     "  ]}}\n}'")
placeholder("PASTE Q1 query result here")
para("Answer: Checkout aborted because the payments service hit "
     "'authorize failed: upstream timeout' from stripe (504, ~5001ms). "
     "Checkout is the symptom; the payment-provider timeout is the cause.")

h2("Q2. Which users were affected?")
code("curl -s 'http://localhost:9200/storefront-logs/_search?pretty' \\\n"
     "  -H 'Content-Type: application/json' -d '{\n"
     "  \"size\": 20, \"_source\": [\"ts\",\"user\",\"trace_id\",\"reason\"],\n"
     "  \"query\": { \"bool\": { \"filter\": [\n"
     "    { \"match\": { \"service\": \"checkout\" } },\n"
     "    { \"match\": { \"level\":   \"ERROR\"    } }\n"
     "  ]}}\n}'")
placeholder("PASTE Q2 query result here")
para("Answer: u-4419 (d4e1), u-4420 (e2a8), u-4422 (aa01). u-4430 succeeded at "
     "14:07 after recovery and was NOT affected.")

h2("Q3. Root cause vs. symptom — correlate by ts and trace_id")
code("curl -s 'http://localhost:9200/storefront-logs/_search?pretty' \\\n"
     "  -H 'Content-Type: application/json' -d '{\n"
     "  \"size\": 20, \"sort\": [ { \"ts\": { \"order\": \"asc\" } } ],\n"
     "  \"query\": { \"match\": { \"trace_id\": \"d4e1\" } }\n}'")
placeholder("PASTE Q3 query result here")
para("Answer: trace_id d4e1 shows payments 'authorize failed: upstream timeout' "
     "(14:05:02.778, ROOT CAUSE) immediately followed by checkout 'checkout "
     "aborted' (14:05:02.901, SYMPTOM). The shared trace_id proves the abort "
     "was downstream of the timeout. Same pattern for e2a8 and aa01 — three "
     "customer-facing failures, one incident.")

# ============================= STEP 5 =============================
h1("6. Step 5 — Ask the AI the Same Questions (ask_logs.py)")
para("Point ask_logs.py at a local model (Ollama) and run the same three "
     "questions. The script PRINTS the AI-generated filter so we can audit it.")
code("export OPENAI_BASE_URL=http://localhost:11434/v1\n"
     "export OPENAI_API_KEY=not-needed\n"
     "export MODEL=llama3.1:8b\n"
     "python3 ask_logs.py sample-logs.jsonl \"why did checkout fail at 14:05?\"\n"
     "python3 ask_logs.py sample-logs.jsonl \"which users were affected?\"\n"
     "python3 ask_logs.py sample-logs.jsonl \"what was the root cause of the 14:05 checkout failures?\"")
para("Paste, for EACH question, the printed '# AI-generated filter' line AND "
     "the lines the script returned:")
placeholder("PASTE ask_logs.py output for Q1 (filter + lines)")
placeholder("PASTE ask_logs.py output for Q2 (filter + lines)")
placeholder("PASTE ask_logs.py output for Q3 (filter + lines)")

# ============================= STEP 6 =============================
h1("7. Step 6 — Verdict (keeping-watch.txt)")
para("The full verdict is in keeping-watch.txt. In short: for Q1/Q2 the AI's "
     "filter overlapped with my by-hand query and found the right lines. For "
     "Q3 it could not correlate across services, so it could not separate root "
     "cause from symptom — I did that by hand using trace_id. I would NOT have "
     "shipped a fix on the AI's answer alone, because it pointed at checkout "
     "(the symptom) and not at the payment-provider timeout (the cause).")

# ============================= EXTRA CREDIT =============================
h1("8. Extra Credit (optional)")
para("Medium tier: AI-workload signals appear in the logs (inference 'kv cache "
     "pressure', gpu_util_pct 97, queue_depth 18) — the box was 'CPU-green but "
     "GPU-saturated.' A symptom-based burn-rate alert for the checkout SLO is "
     "in alert-rules.yml (99.9% over 30 days; alert on burn rate, not raw "
     "error count, to avoid alert fatigue).")
para("Hard tier: the one-page autonomy memo to the IT director is in "
     "HARD-memo.txt.")

# ============================= REPO CONTENTS =============================
h1("9. Repository Contents")
bullet("docker-compose.observability.yml — the four-service stack")
bullet("prometheus.yml — scrape config")
bullet("sample-logs.jsonl — the 14 storefront log lines (the incident)")
bullet("load-logs.sh — bulk-load logs into OpenSearch and verify 14 docs")
bullet("ask_logs.py — natural-language-to-filter bridge (AI proposes, I verify)")
bullet("queries-by-hand.md — the three by-hand OpenSearch queries + answers")
bullet("grafana-queries.md — the p95 + error-ratio PromQL")
bullet("keeping-watch.txt — the incident verdict (AI vs by-hand)")
bullet("agent-log.txt — the AI audit trail (delegated / did / wrong / intervened)")
bullet("alert-rules.yml — SLO burn-rate alert (Medium tier)")
bullet("HARD-memo.txt — autonomy memo to IT director (Hard tier)")
bullet("REPORT.docx — this report")

doc.save("REPORT.docx")
print("Wrote REPORT.docx")
