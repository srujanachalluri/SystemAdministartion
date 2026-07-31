#!/usr/bin/env python3
"""make_report.py - generates REPORT.docx.

Run:  pip install python-docx && python make_report.py
Then open REPORT.docx and paste your real terminal output into the
[PASTE OUTPUT HERE] blocks, and drop in your CI screenshots.
"""
from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()

# Make code blocks readable
styles = doc.styles
body = styles["Normal"]
body.font.name = "Calibri"
body.font.size = Pt(11)


def h(text, level=1):
    doc.add_heading(text, level=level)


def p(text, bold=False):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold
    return par


def code(text):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    return par


def placeholder(label="PASTE OUTPUT HERE"):
    par = doc.add_paragraph()
    run = par.add_run(f"[ {label} ]")
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    return par


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


# =============================================================================
doc.add_heading("Project 13 — A Pipeline for an AI Application", 0)
p("Student: Srujana Challuri")
p("Course: System Administration and Maintenance — Chapter 13, DevOps and SRE")
p("Repository: https://github.com/srujanachalluri/SystemAdministration")
p("Project folder: project-13-ai-pipeline")

# -----------------------------------------------------------------------------
h("1. What I built and why", 1)
p("The school help-desk RAG support bot answers staff questions from a folder "
  "of school policy documents. It was up 100% of the time and it still told a "
  "new teacher the wrong lockdown procedure. No alarm fired, because every "
  "alarm was watching whether the service was UP, and it was always up.")
p("This project adds the reliability apparatus that was missing: a pipeline "
  "that cannot ship a quality regression, a measurable promise about quality "
  "with an error budget behind it, drift monitoring, and a rollback plan "
  "written before the incident instead of during it.")

# -----------------------------------------------------------------------------
h("2. The pipeline and why the stages are in this order", 1)
code("install  ->  lint  ->  unit tests  ->  secret scan  ->  LLM EVAL GATE\n"
     "  free       free        free            free          slow + paid")
p("The LLM eval gate costs real tokens and real minutes. It runs LAST, only on "
  "a change that has already passed every check a human could have caught for "
  "free. Ordering the pipeline for money, not for style.")
p("The workflow lives at .github/workflows/ci.yml at the repository root.")

h("Command — view the pipeline file", 2)
code("cat .github/workflows/ci.yml")
placeholder()

# -----------------------------------------------------------------------------
h("3. The model ID is pinned", 1)
p("Nothing in this project says \"latest\". The exact model ID is pinned in "
  "app/config.py and is enforced by a unit test, so a build fails if anyone "
  "ever replaces it with a floating alias. A pipeline that resolves \"latest\" "
  "at run time is a pipeline that changes without anyone choosing the change.")
code("MODEL_ID = os.environ.get(\"MODEL_ID\", \"gpt-4o-mini-2024-07-18\")")

h("Command — show the RAG bot answering from the fixed corpus", 2)
code("python -m app.rag_service \"Who can release a lockdown?\"")
placeholder()

# -----------------------------------------------------------------------------
h("4. My groundedness threshold, and why it is that number", 1)
p("This is my judgement, not a setting I asked an agent for.", bold=True)

h("The SLI — what I measure", 2)
p("Groundedness rate = grounded answers / total answers, where \"grounded\" "
  "means every factual claim in the answer is supported by the policy text the "
  "bot actually retrieved. An honest refusal counts as grounded. Availability "
  "is deliberately NOT the SLI: the incident happened while the service was at "
  "100% uptime.")

h("The SLO — the promise", 2)
p("99% of answers grounded, over a rolling 30-day window.", bold=True)
p("Why not 99.9%: the readers are adults, every answer is traceable to a named "
  "policy document, and a staff member can and does verify anything unusual "
  "with the IT Help Desk. A wrong answer here is recoverable, not final. "
  "Chasing 99.9% would cost far more in eval spend and human review than the "
  "marginal risk it removes.")
p("Why not 95%: at roughly 20,000 answers a month, 95% means 1,000 staff acting "
  "on fabricated policy — including safety procedure. That is not a number I "
  "am willing to sign my name to.")
p("99% is where the cost of being more careful stops being worth it, and the "
  "cost of being less careful starts landing on a teacher in a real lockdown.")

h("The error budget", 2)
p("100% − 99% = 1%. At 20,000 answers per month that is 200 bad answers I am "
  "permitted to spend. Burn under 75%: ship normally. Burn at or over 75%: "
  "freeze risky deploys. Burn over 100%: stop shipping, roll back to the "
  "last-good model, stabilise.")

h("The CI gate threshold — 0.90, and why it is LOWER than the SLO", 2)
p("The CI eval gate fails the build below 0.90 groundedness. That is "
  "deliberately below the 0.99 production SLO because the eval set is "
  "adversarial: it is weighted toward negative permissions, near-duplicate "
  "numbers, and out-of-corpus questions, which are far rarer in real traffic. "
  "Holding an adversarial set to the production number would either block every "
  "build or push me to soften the eval set — and a softened eval set is worse "
  "than no eval set, because it manufactures false confidence.")

h("Command — show the SLO config and the error budget verdict", 2)
code("cat slo.yml")
placeholder()
code("python monitoring/error_budget.py --slo 0.99 --total 20000 --bad 140")
placeholder()

# -----------------------------------------------------------------------------
h("5. The eval set — 10 cases, chosen for the edges", 1)
p("Quantity is not the point. Each case probes a specific failure mode of a "
  "retrieval system:")
bullet("C01 — happy path: does retrieval work at all.")
bullet("C02 — a specific time (10:00 AM). Hallucinated times are the classic RAG failure.")
bullet("C03 — a procedure location. Must name the real portal, not a plausible one.")
bullet("C04 / C08 — two near-duplicate durations (30 minutes vs 15 minutes) in the "
       "SAME document. Tests whether the bot picks the right one.")
bullet("C05 — a NEGATIVE permission. A helpful bot wants to say yes; policy says no.")
bullet("C06 — the safety-critical lockdown release question. This is the exact "
       "question the bot fabricated in the real incident.")
bullet("C07 — two adjacent-but-different concepts (Secure Building vs Lockdown). "
       "Tests conflation.")
bullet("C09 — a fact in a different document. Tests cross-document routing.")
bullet("C10 — OUT OF CORPUS. The bot must REFUSE, not answer from a loosely "
       "related document. This is the adjacent-but-wrong failure mode.")

h("Command — run the eval gate on the clean build (expect PASS)", 2)
code("python evals/judge.py evals/cases.jsonl \\\n"
     "    --min-groundedness 0.90 --min-pass-rate 0.90\n"
     "echo \"exit code: $?\"")
placeholder()

# -----------------------------------------------------------------------------
h("6. Proof that a regression turns the build RED", 1)
p("A green gate you have never seen go red is proof of nothing. So I regressed "
  "the service on purpose, on branch regression-demo, and watched it fail.")
p("What the regression fabricates:", bold=True)
bullet("A phone number that appears nowhere in the policy docs (555-0142).")
bullet("A 24-hour staff hotline that does not exist.")
bullet("A 20-minute SLA that does not exist.")
bullet("An instruction to release the all clear BY TEXT MESSAGE — which directly "
       "contradicts the lockdown policy's \"never by text message.\"")
p("I chose a regression that contradicts the safety policy specifically. A "
  "regression that only breaks a trivial fact would prove nothing about a bot "
  "that answers questions about lockdowns.")

h("The result — and this is the important part", 2)
p("Lint PASSED. All 7 unit tests PASSED. The secret scan PASSED. "
  "Only the eval gate caught it.", bold=True)
code("groundedness = 1.000  ->  0.100   (threshold 0.90)\n"
     "FAIL: quality regression. Blocking the build.\n"
     "exit code 1  ->  build RED")
p("That separation is the whole lesson. No deterministic test could see this "
  "failure, because the code did not crash and the service did not go down. It "
  "just started being confidently wrong. Only a quality gate can see that.")

h("Command — reproduce the red build", 2)
code("git checkout -b regression-demo\n"
     "# (apply the regressed answer function)\n"
     "ruff check .                  # PASSES\n"
     "pytest tests/unit -q          # PASSES\n"
     "python evals/judge.py evals/cases.jsonl \\\n"
     "    --min-groundedness 0.90 --min-pass-rate 0.90   # FAILS\n"
     "echo \"exit code: $?\"")
placeholder()

h("CI evidence", 2)
p("Failing CI run (red) — link:")
placeholder("PASTE THE FAILING GITHUB ACTIONS RUN URL HERE")
p("Screenshot of the red eval gate:")
placeholder("PASTE SCREENSHOT OF THE FAILED CI RUN HERE")
p("Passing CI run (green) on main — link:")
placeholder("PASTE THE PASSING GITHUB ACTIONS RUN URL HERE")
p("Screenshot of the green pipeline:")
placeholder("PASTE SCREENSHOT OF THE PASSING CI RUN HERE")

# -----------------------------------------------------------------------------
h("7. Least privilege — and why an AI agent must not have write access", 1)
code("permissions:\n  contents: read")
p("The CI job can read the repository and nothing else. It cannot push to main, "
  "cannot open or merge a pull request, and cannot edit its own workflow file.")
p("Why that matters here specifically:", bold=True)
p("A coding agent wrote a large part of the code that runs in this job. An "
  "agent with write access to a protected branch could modify the very gate "
  "that is supposed to judge its work — lower the threshold, delete a failing "
  "eval case, or push straight past the pull request that triggers the gate. "
  "The gate would still be green. It would just be green about nothing.")
p("The gate only means something if the thing being judged cannot reach the "
  "judge. Read-only CI permissions plus branch protection on main is what "
  "keeps that true. The agent proposes; a human merges.")

# -----------------------------------------------------------------------------
h("8. Drift monitoring — and what it cannot see", 1)
p("PSI (Population Stability Index) compares the reference window from when the "
  "model was last validated against the current window. Under 0.10 is stable, "
  "0.10–0.20 is worth investigating, over 0.20 means the model may be operating "
  "off-distribution.")

h("Command — run the drift check", 2)
code("python monitoring/psi.py monitoring/data/reference.txt \\\n"
     "                        monitoring/data/current.txt")
placeholder()

h("The honest limitation", 2)
p("PSI detects DATA drift only. It is blind to CONCEPT drift.", bold=True)
p("Here is exactly how that bites this service. When the policy documents were "
  "updated for the 2026 school year, the questions staff asked did not change "
  "at all — same wording, same volume, same length distribution. PSI would have "
  "stayed completely flat while every answer quietly went stale against the new "
  "policy. My monitoring would have said everything was fine.")
p("What actually catches that: a corpus version check that re-runs the eval set "
  "whenever the policy documents change, and a human spot-reviewing sampled "
  "answers. A metric cannot notice that the world moved. A person has to look.")

# -----------------------------------------------------------------------------
h("9. Rollback vs retrain — decided in advance", 1)
p("Roll back immediately if:", bold=True)
bullet("The eval gate fails on main after a deploy.")
bullet("Groundedness drops below 0.95 in any 24-hour window.")
bullet("Any staff member reports a fabricated safety-procedure answer.")
p("Retrain or re-index instead if:", bold=True)
bullet("Groundedness is stable but the answers are WRONG because the source "
       "policy documents changed.")
bullet("PSI stays above 0.20 for seven days with no code change.")
p("Rollback is the reflex, retrain is the decision. Roll back first, diagnose "
  "second. Retraining under incident pressure is how you ship a second bug on "
  "top of the first one.")

# -----------------------------------------------------------------------------
h("10. Where the agent stopped and I started", 1)
p("Full detail is in agent-log.txt. In short: I let the agent write the "
  "pipeline YAML, the glue code, and the test scaffolding. I did not delegate "
  "the SLO number, the eval cases, the groundedness threshold, or the "
  "rollback policy.")
p("One concrete example: the agent's first offline judge scored groundedness by "
  "plain word overlap. I tested it against a fabricated answer and it passed — "
  "invented phone numbers and durations are short tokens that barely move an "
  "overlap ratio. I specified the extra rule myself: every token containing a "
  "digit must appear verbatim in the retrieved context, no exceptions. That is "
  "what actually catches \"call 555-0142\" and \"within 20 minutes.\"")
p("The agent is fast at building the apparatus. It has no standing to decide "
  "how often this service is allowed to be wrong about a lockdown.")

doc.save("REPORT.docx")
print("Wrote REPORT.docx")
